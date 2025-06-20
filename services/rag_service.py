"""
RAG (Retrieval-Augmented Generation) Service
Handles document ingestion, vector storage, similarity search, and context retrieval
"""

import os
import json
import hashlib
import asyncio
from typing import Dict, List, Optional, Tuple, Union
from dataclasses import dataclass, asdict
from datetime import datetime
import logging
from pathlib import Path
import numpy as np
from collections import defaultdict

logger = logging.getLogger(__name__)

try:
    import chromadb
    from chromadb.config import Settings
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False
    logger.warning("ChromaDB not available. Install with: pip install chromadb")

try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False
    logger.warning("SentenceTransformers not available. Install with: pip install sentence-transformers")

try:
    import PyPDF2
    import docx
    from bs4 import BeautifulSoup
    DOCUMENT_PARSERS_AVAILABLE = True
except ImportError:
    DOCUMENT_PARSERS_AVAILABLE = False
    logger.warning("Document parsers not available. Install with: pip install PyPDF2 python-docx beautifulsoup4")


@dataclass
class Document:
    """Document structure for RAG"""
    id: str
    content: str
    metadata: Dict
    embedding: Optional[List[float]] = None
    created_at: datetime = None
    
    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.now()


@dataclass
class SearchResult:
    """Search result structure"""
    document: Document
    score: float
    rank: int


@dataclass
class RAGResponse:
    """RAG-enhanced response structure"""
    answer: str
    sources: List[SearchResult]
    context_used: str
    confidence: float
    model_used: str
    retrieval_time: float
    generation_time: float


class DocumentProcessor:
    """Process various document formats"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        if not DOCUMENT_PARSERS_AVAILABLE:
            raise ImportError("Document parsers not available")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DOCUMENT_PARSERS_AVAILABLE:
            raise ImportError("Document parsers not available")
        
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_html(file_path: str) -> str:
        """Extract text from HTML file"""
        if not DOCUMENT_PARSERS_AVAILABLE:
            raise ImportError("Document parsers not available")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                return soup.get_text().strip()
        except Exception as e:
            logger.error(f"Error extracting text from HTML {file_path}: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            return ""
    
    @classmethod
    def process_file(cls, file_path: str) -> str:
        """Process file based on extension"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        processors = {
            '.pdf': cls.extract_text_from_pdf,
            '.docx': cls.extract_text_from_docx,
            '.doc': cls.extract_text_from_docx,
            '.html': cls.extract_text_from_html,
            '.htm': cls.extract_text_from_html,
            '.txt': cls.extract_text_from_txt,
            '.md': cls.extract_text_from_txt,
        }
        
        processor = processors.get(extension)
        if processor:
            return processor(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {extension}")


class TextChunker:
    """Split text into manageable chunks"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def chunk_text(self, text: str, metadata: Dict = None) -> List[Document]:
        """Split text into chunks with overlap"""
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + self.chunk_size - 100:
                    end = sentence_end + 1
                else:
                    # Look for paragraph breaks
                    para_break = text.rfind('\n\n', start, end)
                    if para_break > start + self.chunk_size - 200:
                        end = para_break + 2
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                doc_id = hashlib.md5(f"{metadata.get('source', 'unknown')}_{chunk_id}".encode()).hexdigest()
                
                chunk_metadata = {
                    **(metadata or {}),
                    'chunk_id': chunk_id,
                    'start_pos': start,
                    'end_pos': end,
                    'chunk_size': len(chunk_text)
                }
                
                chunks.append(Document(
                    id=doc_id,
                    content=chunk_text,
                    metadata=chunk_metadata
                ))
                
                chunk_id += 1
            
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks


class EmbeddingService:
    """Handle text embeddings"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.model_name = model_name
        self.model = None
        self._load_model()
    
    def _load_model(self):
        """Load the embedding model"""
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            logger.warning("SentenceTransformers not available, using fallback embeddings")
            return
        
        try:
            self.model = SentenceTransformer(self.model_name)
            logger.info(f"Loaded embedding model: {self.model_name}")
        except Exception as e:
            logger.error(f"Failed to load embedding model: {e}")
            self.model = None
    
    def embed_text(self, text: str) -> List[float]:
        """Generate embedding for text"""
        if self.model is None:
            # Fallback: simple hash-based embedding
            return self._fallback_embedding(text)
        
        try:
            embedding = self.model.encode(text, convert_to_tensor=False)
            return embedding.tolist()
        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            return self._fallback_embedding(text)
    
    def embed_batch(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for multiple texts"""
        if self.model is None:
            return [self._fallback_embedding(text) for text in texts]
        
        try:
            embeddings = self.model.encode(texts, convert_to_tensor=False)
            return embeddings.tolist()
        except Exception as e:
            logger.error(f"Error generating batch embeddings: {e}")
            return [self._fallback_embedding(text) for text in texts]
    
    def _fallback_embedding(self, text: str, dim: int = 384) -> List[float]:
        """Simple fallback embedding based on text hash"""
        # Create a deterministic embedding based on text hash
        hash_obj = hashlib.sha256(text.encode())
        hash_bytes = hash_obj.digest()
        
        # Convert to float array
        embedding = []
        for i in range(0, min(len(hash_bytes), dim // 8)):
            byte_chunk = hash_bytes[i:i+8] if i+8 <= len(hash_bytes) else hash_bytes[i:] + b'\x00' * (8 - (len(hash_bytes) - i))
            value = int.from_bytes(byte_chunk, byteorder='big', signed=False)
            normalized = (value / (2**64 - 1)) * 2 - 1  # Normalize to [-1, 1]
            embedding.append(normalized)
        
        # Pad or truncate to desired dimension
        while len(embedding) < dim:
            embedding.append(0.0)
        
        return embedding[:dim]


class VectorStore:
    """Vector storage and similarity search"""
    
    def __init__(self, collection_name: str = "rag_documents", persist_directory: str = "./vector_db"):
        self.collection_name = collection_name
        self.persist_directory = persist_directory
        self.client = None
        self.collection = None
        self._initialize_store()
    
    def _initialize_store(self):
        """Initialize the vector store"""
        if not CHROMADB_AVAILABLE:
            logger.warning("ChromaDB not available, using in-memory fallback")
            self.documents = {}
            self.embeddings = {}
            return
        
        try:
            # Create persist directory
            os.makedirs(self.persist_directory, exist_ok=True)
            
            # Initialize ChromaDB client
            self.client = chromadb.PersistentClient(path=self.persist_directory)
            
            # Get or create collection
            try:
                self.collection = self.client.get_collection(name=self.collection_name)
            except ValueError:
                self.collection = self.client.create_collection(
                    name=self.collection_name,
                    metadata={"hnsw:space": "cosine"}
                )
            
            logger.info(f"Initialized vector store: {self.collection_name}")
            
        except Exception as e:
            logger.error(f"Failed to initialize ChromaDB: {e}")
            logger.warning("Falling back to in-memory storage")
            self.client = None
            self.collection = None
            self.documents = {}
            self.embeddings = {}
    
    def add_documents(self, documents: List[Document]) -> bool:
        """Add documents to the vector store"""
        if self.collection is None:
            return self._add_documents_fallback(documents)
        
        try:
            ids = [doc.id for doc in documents]
            embeddings = [doc.embedding for doc in documents if doc.embedding]
            metadatas = [doc.metadata for doc in documents]
            documents_text = [doc.content for doc in documents]
            
            if not embeddings:
                logger.warning("No embeddings found in documents")
                return False
            
            self.collection.add(
                ids=ids,
                embeddings=embeddings,
                metadatas=metadatas,
                documents=documents_text
            )
            
            logger.info(f"Added {len(documents)} documents to vector store")
            return True
            
        except Exception as e:
            logger.error(f"Error adding documents to vector store: {e}")
            return False
    
    def _add_documents_fallback(self, documents: List[Document]) -> bool:
        """Fallback method for adding documents"""
        try:
            for doc in documents:
                self.documents[doc.id] = doc
                if doc.embedding:
                    self.embeddings[doc.id] = doc.embedding
            return True
        except Exception as e:
            logger.error(f"Error in fallback document storage: {e}")
            return False
    
    def search(self, query_embedding: List[float], n_results: int = 5) -> List[SearchResult]:
        """Search for similar documents"""
        if self.collection is None:
            return self._search_fallback(query_embedding, n_results)
        
        try:
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results
            )
            
            search_results = []
            for i, (doc_id, distance, metadata, content) in enumerate(zip(
                results['ids'][0],
                results['distances'][0],
                results['metadatas'][0],
                results['documents'][0]
            )):
                # Convert distance to similarity score
                score = 1.0 - distance
                
                document = Document(
                    id=doc_id,
                    content=content,
                    metadata=metadata
                )
                
                search_results.append(SearchResult(
                    document=document,
                    score=score,
                    rank=i + 1
                ))
            
            return search_results
            
        except Exception as e:
            logger.error(f"Error searching vector store: {e}")
            return []
    
    def _search_fallback(self, query_embedding: List[float], n_results: int = 5) -> List[SearchResult]:
        """Fallback search using cosine similarity"""
        try:
            if not self.embeddings:
                return []
            
            similarities = []
            query_array = np.array(query_embedding)
            
            for doc_id, embedding in self.embeddings.items():
                doc_array = np.array(embedding)
                
                # Cosine similarity
                dot_product = np.dot(query_array, doc_array)
                norm_query = np.linalg.norm(query_array)
                norm_doc = np.linalg.norm(doc_array)
                
                if norm_query > 0 and norm_doc > 0:
                    similarity = dot_product / (norm_query * norm_doc)
                else:
                    similarity = 0.0
                
                similarities.append((doc_id, similarity))
            
            # Sort by similarity and take top results
            similarities.sort(key=lambda x: x[1], reverse=True)
            top_results = similarities[:n_results]
            
            search_results = []
            for rank, (doc_id, score) in enumerate(top_results, 1):
                if doc_id in self.documents:
                    search_results.append(SearchResult(
                        document=self.documents[doc_id],
                        score=score,
                        rank=rank
                    ))
            
            return search_results
            
        except Exception as e:
            logger.error(f"Error in fallback search: {e}")
            return []
    
    def get_collection_stats(self) -> Dict:
        """Get statistics about the collection"""
        if self.collection is None:
            return {
                "total_documents": len(self.documents),
                "storage_type": "in_memory_fallback"
            }
        
        try:
            count = self.collection.count()
            return {
                "total_documents": count,
                "collection_name": self.collection_name,
                "storage_type": "chromadb"
            }
        except Exception as e:
            logger.error(f"Error getting collection stats: {e}")
            return {"error": str(e)}


class RAGService:
    """Main RAG service orchestrator"""
    
    def __init__(self, 
                 collection_name: str = "rag_documents",
                 embedding_model: str = "all-MiniLM-L6-v2",
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200):
        
        self.embedding_service = EmbeddingService(embedding_model)
        self.text_chunker = TextChunker(chunk_size, chunk_overlap)
        self.vector_store = VectorStore(collection_name)
        self.document_processor = DocumentProcessor()
        
        # RAG configuration
        self.max_context_length = 4000
        self.min_similarity_threshold = 0.3
        self.max_sources = 5
        
        # Statistics
        self.stats = defaultdict(int)
    
    def ingest_file(self, file_path: str, metadata: Dict = None) -> bool:
        """Ingest a file into the RAG system"""
        try:
            # Extract text from file
            text = self.document_processor.process_file(file_path)
            if not text:
                logger.warning(f"No text extracted from {file_path}")
                return False
            
            # Prepare metadata
            file_metadata = {
                "source": str(file_path),
                "file_name": Path(file_path).name,
                "file_size": os.path.getsize(file_path),
                "ingested_at": datetime.now().isoformat(),
                **(metadata or {})
            }
            
            return self.ingest_text(text, file_metadata)
            
        except Exception as e:
            logger.error(f"Error ingesting file {file_path}: {e}")
            return False
    
    def ingest_text(self, text: str, metadata: Dict = None) -> bool:
        """Ingest text into the RAG system"""
        try:
            # Chunk the text
            chunks = self.text_chunker.chunk_text(text, metadata)
            if not chunks:
                logger.warning("No chunks created from text")
                return False
            
            # Generate embeddings
            texts = [chunk.content for chunk in chunks]
            embeddings = self.embedding_service.embed_batch(texts)
            
            # Add embeddings to chunks
            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = embedding
            
            # Store in vector database
            success = self.vector_store.add_documents(chunks)
            
            if success:
                self.stats['documents_ingested'] += len(chunks)
                logger.info(f"Successfully ingested {len(chunks)} chunks")
            
            return success
            
        except Exception as e:
            logger.error(f"Error ingesting text: {e}")
            return False
    
    def search_documents(self, query: str, n_results: int = 5) -> List[SearchResult]:
        """Search for relevant documents"""
        try:
            # Generate query embedding
            query_embedding = self.embedding_service.embed_text(query)
            
            # Search vector store
            results = self.vector_store.search(query_embedding, n_results)
            
            # Filter by similarity threshold
            filtered_results = [
                result for result in results 
                if result.score >= self.min_similarity_threshold
            ]
            
            self.stats['searches_performed'] += 1
            return filtered_results
            
        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []
    
    def generate_rag_response(self, 
                            query: str, 
                            model_name: str = "llama2:7b-chat",
                            max_sources: int = None) -> RAGResponse:
        """Generate RAG-enhanced response"""
        from services.ollama_service import ollama_service
        import time
        
        start_time = time.time()
        
        try:
            # Search for relevant documents
            search_start = time.time()
            max_sources = max_sources or self.max_sources
            search_results = self.search_documents(query, max_sources)
            retrieval_time = time.time() - search_start
            
            # Build context from search results
            context_parts = []
            total_context_length = 0
            
            for result in search_results:
                content = result.document.content
                if total_context_length + len(content) <= self.max_context_length:
                    source_info = result.document.metadata.get('source', 'Unknown')
                    context_parts.append(f"[Source: {source_info}]\n{content}")
                    total_context_length += len(content)
                else:
                    break
            
            context = "\n\n".join(context_parts)
            
            # Create RAG prompt
            rag_prompt = self._build_rag_prompt(query, context)
            
            # Generate response using Ollama
            generation_start = time.time()
            
            # Use the chat service for generation
            from services.chat_service import chat_service
            
            # Create a temporary session for RAG
            rag_messages = [
                {"role": "system", "content": "You are a helpful assistant that answers questions based on provided context. Always cite your sources."},
                {"role": "user", "content": rag_prompt}
            ]
            
            response_content = ""
            for response_chunk in ollama_service.chat(
                model=model_name,
                messages=rag_messages,
                stream=False
            ):
                response_content = response_chunk.content
                break  # Take the first (complete) response
            
            generation_time = time.time() - generation_start
            
            # Calculate confidence based on search results
            confidence = self._calculate_confidence(search_results, query)
            
            # Update statistics
            self.stats['rag_responses_generated'] += 1
            
            return RAGResponse(
                answer=response_content,
                sources=search_results,
                context_used=context,
                confidence=confidence,
                model_used=model_name,
                retrieval_time=retrieval_time,
                generation_time=generation_time
            )
            
        except Exception as e:
            logger.error(f"Error generating RAG response: {e}")
            return RAGResponse(
                answer=f"Error generating response: {str(e)}",
                sources=[],
                context_used="",
                confidence=0.0,
                model_used=model_name,
                retrieval_time=0.0,
                generation_time=time.time() - start_time
            )
    
    def _build_rag_prompt(self, query: str, context: str) -> str:
        """Build the RAG prompt"""
        if not context:
            return f"""Question: {query}

I don't have any relevant context to answer this question. Please provide a general response based on your knowledge, but mention that you don't have specific context for this query."""
        
        return f"""Context:
{context}

Question: {query}

Please answer the question based on the provided context. If the context doesn't contain enough information to answer the question, say so. Always cite the sources mentioned in the context."""
    
    def _calculate_confidence(self, search_results: List[SearchResult], query: str) -> float:
        """Calculate confidence score for the response"""
        if not search_results:
            return 0.0
        
        # Base confidence on average similarity score
        avg_score = sum(result.score for result in search_results) / len(search_results)
        
        # Boost confidence if we have multiple relevant sources
        source_bonus = min(len(search_results) * 0.1, 0.3)
        
        # Penalize if top result has low similarity
        top_score_penalty = 0.0
        if search_results and search_results[0].score < 0.5:
            top_score_penalty = 0.2
        
        confidence = min(avg_score + source_bonus - top_score_penalty, 1.0)
        return max(confidence, 0.0)
    
    def get_stats(self) -> Dict:
        """Get RAG service statistics"""
        vector_stats = self.vector_store.get_collection_stats()
        
        return {
            "rag_stats": dict(self.stats),
            "vector_store": vector_stats,
            "configuration": {
                "chunk_size": self.text_chunker.chunk_size,
                "chunk_overlap": self.text_chunker.chunk_overlap,
                "max_context_length": self.max_context_length,
                "min_similarity_threshold": self.min_similarity_threshold,
                "max_sources": self.max_sources,
                "embedding_model": self.embedding_service.model_name
            }
        }
    
    def clear_documents(self) -> bool:
        """Clear all documents from the vector store"""
        try:
            if self.vector_store.collection:
                # ChromaDB doesn't have a clear method, so we delete and recreate
                self.vector_store.client.delete_collection(self.vector_store.collection_name)
                self.vector_store.collection = self.vector_store.client.create_collection(
                    name=self.vector_store.collection_name,
                    metadata={"hnsw:space": "cosine"}
                )
            else:
                # Fallback storage
                self.vector_store.documents.clear()
                self.vector_store.embeddings.clear()
            
            self.stats['documents_cleared'] += 1
            logger.info("Cleared all documents from vector store")
            return True
            
        except Exception as e:
            logger.error(f"Error clearing documents: {e}")
            return False


# Global RAG service instance
rag_service = RAGService()